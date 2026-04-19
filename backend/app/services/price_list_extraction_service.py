"""Extract text from uploaded price list files."""
import logging

logger = logging.getLogger(__name__)


def extract_text_from_file(
    file_content: bytes,
    file_type: str,
    *,
    company_id: str | None = None,
    import_id: str | None = None,
) -> str:
    """Extract text from a price list file.
    Supports: excel, csv, pdf, word.
    Returns plain text representation.

    company_id / import_id are optional — they're threaded to the Intelligence
    layer when PDF vision fallback runs (Phase 2c-1). Non-PDF types ignore them.
    """
    if file_type in ("excel", "xlsx", "xls"):
        return _extract_excel(file_content)
    elif file_type == "csv":
        return _extract_csv(file_content)
    elif file_type == "pdf":
        return _extract_pdf(file_content, company_id=company_id, import_id=import_id)
    elif file_type in ("word", "docx", "doc"):
        return _extract_word(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def _extract_excel(content: bytes) -> str:
    """Extract text from Excel using openpyxl."""
    import io
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            vals = [str(c) if c is not None else "" for c in row]
            if any(v.strip() for v in vals):
                lines.append(" | ".join(vals))
    wb.close()
    return "\n".join(lines)


def _extract_csv(content: bytes) -> str:
    """Extract text from CSV."""
    import csv
    import io

    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    lines: list[str] = []
    for row in reader:
        if any(c.strip() for c in row):
            lines.append(" | ".join(row))
    return "\n".join(lines)


def _extract_pdf(
    content: bytes,
    *,
    company_id: str | None = None,
    import_id: str | None = None,
) -> str:
    """Extract text from PDF.

    Strategy 1: pdfplumber (fast, handles text-layer PDFs).
    Strategy 2: Claude vision API (fallback for scanned/image-only PDFs).
    """
    # Strategy 1 — pdfplumber
    try:
        import io

        import pdfplumber

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            lines: list[str] = []
            for page in pdf.pages[:20]:
                text = page.extract_text()
                if text:
                    lines.append(text)
            text_result = "\n".join(lines)
            if text_result.strip():
                return text_result
            logger.info("pdfplumber extracted no text — falling back to Claude vision")
    except ImportError:
        logger.warning("pdfplumber not installed — trying Claude vision fallback")
    except Exception as e:
        logger.warning("pdfplumber failed (%s) — trying Claude vision fallback", e)

    # Strategy 2 — Claude vision (handles scanned/image-based PDFs)
    return _extract_pdf_via_claude(content, company_id=company_id, import_id=import_id)


def _extract_pdf_via_claude(
    content: bytes,
    *,
    db=None,
    company_id: str | None = None,
    import_id: str | None = None,
) -> str:
    """Send the PDF to Claude via the managed `pricing.extract_pdf_text` prompt.

    Phase 2c-1 migration — the raw PDF is redacted from the audit row (only
    sha256 + bytes length are preserved). company_id is required to scope the
    execution row; if the caller can't provide it yet we fall back to an
    unscoped audit row (caller_module alone).

    import_id is populated when the PriceListImport row already exists at call
    time (future refactor). Today, extraction runs BEFORE the import record
    is created, so import_id is typically None — the subsequent analyze step
    carries that linkage.
    """
    import base64

    from app.config import settings

    if not settings.ANTHROPIC_API_KEY:
        logger.warning("No ANTHROPIC_API_KEY — cannot use Claude PDF fallback")
        return ""

    if db is None:
        # The managed layer requires a Session. Open a short-lived one so the
        # caller can keep the simple synchronous API (extract_text_from_file
        # currently doesn't take a db). Close before returning.
        from app.database import SessionLocal

        local_db = SessionLocal()
        try:
            return _extract_pdf_via_claude(
                content, db=local_db, company_id=company_id, import_id=import_id
            )
        finally:
            local_db.close()

    try:
        from app.services.intelligence import intelligence_service

        b64 = base64.standard_b64encode(content).decode("utf-8")
        result = intelligence_service.execute(
            db,
            prompt_key="pricing.extract_pdf_text",
            variables={},
            company_id=company_id,
            caller_module="price_list_extraction_service._extract_pdf_via_claude",
            caller_entity_type="price_list_import" if import_id else None,
            caller_entity_id=import_id,
            caller_price_list_import_id=import_id,
            content_blocks=[
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": b64,
                    },
                }
            ],
        )
        extracted = result.response_text or ""
        logger.info(
            "Claude PDF extraction returned %d chars (tokens: in=%s out=%s)",
            len(extracted),
            result.input_tokens,
            result.output_tokens,
        )
        return extracted
    except Exception as e:
        logger.error("Claude PDF extraction failed: %s", e)
        return ""


def _extract_word(content: bytes) -> str:
    """Extract text from Word docx."""
    try:
        import io

        from docx import Document

        doc = Document(io.BytesIO(content))
        lines: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except ImportError:
        logger.warning("python-docx not installed — Word extraction unavailable")
        return ""
